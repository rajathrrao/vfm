from docx import Document
from docx.shared import Inches, Pt

def digi_sign_doc(file_id, dept, status, comment, digital_sign):
    file_name = file_id+'_digital_signatures.docx'

    if status=='composed':
        document=Document()
        document.add_heading(file_id, 0)
        table=document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.style.font.size = Pt(13)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('Department').bold=True
        hdr_cells[1].paragraphs[0].add_run('Status').bold=True
        hdr_cells[2].paragraphs[0].add_run('Comment').bold=True
        hdr_cells[3].paragraphs[0].add_run('Digital Signature').bold=True
    else:
        try:
            document=Document(file_name)
        except:
            print('Error while opening document')
            return False
        table=document.tables[0]

    row_cells = table.add_row().cells
    row_cells[0].text = dept
    row_cells[1].text = status
    row_cells[2].text = comment
    run = row_cells[3].paragraphs[0].add_run()
    try:
        run.add_picture(digital_sign, width = Inches(1.25), height = Inches(1))
    except:
        print('Error while adding image')
        return False
    try:
        document.save(file_name)
        return True
    except:
        print('Error while saving document')
        return False
